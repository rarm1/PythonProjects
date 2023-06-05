import os
import time
from pathlib import Path

import comtypes.client
import docx

from AOL.Resources.agency import Agency
from AOL.Resources.constants import Constants
from AOL.Resources import constants as const
from AOL.Resources.email_creation import Email
from AOL.Resources.fund_details import FundDetails
from AOL.Resources.output_document import OutputDocument


def allowable_execution(row: int) -> object:
	const = Constants()
	if const.FundNameColumn is None:
		return False
	if const.Reader_Sheet.cell(row=row, column=const.FundNameColumn).value is None:
		return False
	else:
		return True


def create_new_document(filepath, fund, agent):
	search = []
	doc_to_read = docx.Document(filepath)
	for paragraph in doc_to_read.paragraphs:
		for idx, r in enumerate(paragraph.runs):
			replace = False
			if "{" in r.text:
				search = []
				for letter in r.text:
					if letter == "{":
						replace = True
					if replace:
						search.append(letter)
					if letter == "}":
						break
				search = "".join(search)
			if replace:
				replacement_text = str
				to_search = search[1:-1]
				if "date" in to_search.lower():
					replacement_text = r.text.replace(search, str(const.date_format))
				elif "fund_name" in to_search.lower():
					replacement_text = r.text.replace(search, fund.Fund_Name)
				elif "isin" in to_search.lower():
					replacement_text = r.text.replace(search, fund.ISIN)
				elif "designation_list" in to_search.lower():
					replacement_text = r.text.replace(search, fund.Designation_List)
				
				elif "addr" in to_search.lower():
					replacement_text = r.text.replace(search, fund.BNYAddress)
				elif "agency_code" in to_search.lower():
					if agent.Agency_Code is None:
						replacement_text = r.text.replace(search, "")
					else:
						replacement_text = r.text.replace(search, agent.Agency_Code)
				paragraph.runs[idx].text = replacement_text
	filepath = Path(filepath)
	doc_to_read.save(filepath)


def process_fund(row, agent):
	if allowable_execution(row):
		print(f"Iteration: {row - 1}")
		fund = FundDetails(row)
		out = OutputDocument()
		out.AOL_Filename = f'{fund.Fund_Name} AOL.docx'
		out.DA_Filename = f'{fund.Fund_Name} DA.docx'
		current_path = str(Path().parent.resolve())
		current_path = str(current_path)
		current_path += "\\Output"
		out.Location = Path(current_path, fund.Fund_Name)
		if not os.path.exists(out.Location):
			os.makedirs(out.Location)
		out.create_files()
		out.AOL_pdf_filename = str(Path(out.AOL_Written_File_Path).with_suffix(".pdf"))
		out.DA_pdf_filename = str(Path(out.DA_Written_File_Path).with_suffix(".pdf"))
		create_new_document(out.AOL_Written_File_Path, fund, agent)
		create_new_document(out.DA_Written_File_Path, fund, agent)
		word = comtypes.client.CreateObject('Word.Application')
		word.Visible = False
		doc1 = word.Documents.Open(os.path.abspath(out.AOL_Written_File_Path))
		doc1.SaveAs(out.AOL_pdf_filename, 17)
		doc2 = word.Documents.Open(os.path.abspath(out.DA_Written_File_Path))
		doc2.SaveAs(out.DA_pdf_filename, 17)
		word.Quit()
		email = Email(agent, fund)
		email.save_email(out)


def main():
	input_document = Constants()
	input_sheet = input_document.Reader_Sheet
	for row in range(2, input_sheet.max_row + 1):
		agent = Agency(row)
		if agent.Agent_Phone is not None:
			process_fund(row, agent)


if __name__ == '__main__':
	start_time = time.time()
	
	main()
	
	end_time = time.time()
	total_time = end_time - start_time
	
	print(f"Program completed in {total_time} seconds.")
